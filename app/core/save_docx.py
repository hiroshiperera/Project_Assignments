from docx import Document
import os

class Save_Dcoument:
    def __init__(self):
        self.file_path = "D:/DataScience/Assignments/Rag_Project/app/output/output.docx"
        # Create new doc only if it doesn't exist
        if not os.path.exists(self.file_path):
            self.doc = Document()
            self.doc.save(self.file_path)



    def save_response_to_docx(self, title, content):
        # Open existing document
        doc = Document(self.file_path)
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)
        doc.save(self.file_path)
